"""
converter.py — Multi-format document export
Converts markdown text to .txt / .md / .docx / .xlsx / .pdf bytes.
All functions return bytes. Errors raise exceptions — callers must handle.
"""
import re
from io import BytesIO

SUPPORTED_FORMATS = {'md', 'txt', 'docx', 'xlsx', 'pdf'}


# ─── Public API ──────────────────────────────────────────────────────────────

def convert(text: str, fmt: str) -> bytes:
    """Convert markdown text to target format. Returns bytes."""
    fmt = fmt.lower().strip()
    if fmt == 'txt':  return to_txt(text)
    if fmt == 'docx': return to_docx(text)
    if fmt == 'xlsx': return to_xlsx(text)
    if fmt == 'pdf':  return to_pdf(text)
    return to_md(text)  # default: .md


# ─── Plain text ───────────────────────────────────────────────────────────────

def to_txt(text: str) -> bytes:
    return text.encode('utf-8')


def to_md(text: str) -> bytes:
    return text.encode('utf-8')


# ─── Word document (.docx) ────────────────────────────────────────────────────

def to_docx(text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()

    # Set default font to support Thai
    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = Pt(14)

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('# '):
            doc.add_heading(_strip_inline(line[2:].strip()), level=1)
        elif line.startswith('## '):
            doc.add_heading(_strip_inline(line[3:].strip()), level=2)
        elif line.startswith('### '):
            doc.add_heading(_strip_inline(line[4:].strip()), level=3)

        # Table — collect all consecutive | rows
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows like |---|---|
                if not re.match(r'^\|[\s\-:|]+\|$', row_text):
                    table_lines.append(row_text)
                i += 1

            if table_lines:
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    rows.append(cells)
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for ci in range(max_cols):
                        cell_text = row[ci] if ci < len(row) else ''
                        cell = table.cell(ri, ci)
                        cell.text = _strip_inline(cell_text)
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue  # i already advanced in inner loop

        # List bullet
        elif re.match(r'^[-*]\s', line):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, line[2:].strip())

        # Horizontal rule
        elif re.match(r'^[-*_]{3,}\s*$', line.strip()):
            doc.add_paragraph('─' * 50)

        # Empty line → skip
        elif line.strip() == '':
            pass

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line.strip())

        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_inline(text: str) -> str:
    """Remove markdown inline markers for plain text contexts."""
    prev = None
    while prev != text:
        prev = text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def _add_inline_runs(paragraph, text: str):
    """Add paragraph runs, converting **bold** to bold runs."""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for j, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            run.bold = (j % 2 == 1)


# ─── Excel spreadsheet (.xlsx) ────────────────────────────────────────────────

def to_xlsx(text: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Document'

    lines = text.split('\n')

    # Find first markdown table
    table_lines = []
    in_table = False
    for line in lines:
        if line.strip().startswith('|'):
            in_table = True
            # Skip separator rows |---|---|
            if not re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                table_lines.append(line.strip())
        elif in_table:
            break  # stop at first gap after table

    if table_lines:
        for row_idx, tl in enumerate(table_lines, 1):
            cells = [c.strip() for c in tl.strip('|').split('|')]
            for col_idx, cell_text in enumerate(cells, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=cell_text)
                if row_idx == 1:
                    cell.font = Font(bold=True, color='FFFFFF')
                    cell.fill = PatternFill('solid', fgColor='1B2130')
                cell.alignment = Alignment(wrap_text=True)

        # Auto column width
        for col in ws.columns:
            max_len = max((len(str(c.value or '')) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    else:
        # No table found — put full text in A1 with a note
        ws['A1'] = 'หมายเหตุ: ไม่พบตารางในเอกสาร — แสดงเนื้อหาทั้งหมดแทน'
        ws['A1'].font = Font(italic=True, color='888888')
        ws['A2'] = text
        ws['A2'].alignment = Alignment(wrap_text=True)
        ws.column_dimensions['A'].width = 100
        ws.row_dimensions[2].height = 400

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── PDF (.pdf) ───────────────────────────────────────────────────────────────

def to_pdf(text: str) -> bytes:
    try:
        _MAX_PDF_CHARS = 100_000  # ~100K characters
        if len(text) > _MAX_PDF_CHARS:
            text = text[:_MAX_PDF_CHARS] + '\n\n[⚠️ เอกสารถูกตัดเนื่องจากมีขนาดใหญ่เกินไป]'
        import markdown as md_lib
        from weasyprint import HTML
    except ImportError as e:
        raise RuntimeError(f"PDF export ต้องการ weasyprint และ markdown: {e}")

    try:
        html_body = md_lib.markdown(
            text,
            extensions=['tables', 'fenced_code', 'nl2br']
        )

        html = f"""<!DOCTYPE html>
<html lang="th">
<head>
<meta charset="UTF-8">
<style>
  body {{
    font-family: 'Norasi', 'TH Sarabun New', 'Garuda', 'Loma', sans-serif;
    font-size: 14pt;
    line-height: 1.9;
    margin: 0;
    color: #1a1a1a;
  }}
  @page {{
    margin: 2cm 2.5cm;
  }}
  h1 {{ font-size: 20pt; margin: 0 0 10pt; border-bottom: 2px solid #333; padding-bottom: 6pt; }}
  h2 {{ font-size: 17pt; margin: 16pt 0 8pt; }}
  h3 {{ font-size: 15pt; margin: 12pt 0 6pt; }}
  p  {{ margin: 6pt 0; }}
  ul, ol {{ padding-left: 20pt; margin: 6pt 0; }}
  li {{ margin: 3pt 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12pt 0; }}
  th, td {{ border: 1px solid #aaa; padding: 6pt 10pt; text-align: left; }}
  th {{ background: #e8eaf0; font-weight: bold; }}
  tr:nth-child(even) td {{ background: #f8f9fb; }}
  hr {{ border: none; border-top: 1px solid #ccc; margin: 14pt 0; }}
  code {{ background: #f0f0f0; padding: 1pt 4pt; border-radius: 3pt; font-size: 11pt; }}
  pre  {{ background: #f0f0f0; padding: 10pt; border-radius: 4pt; overflow-x: auto; }}
  blockquote {{ border-left: 4px solid #aaa; padding-left: 12pt; color: #555; margin: 10pt 0; }}
  strong {{ font-weight: bold; }}
</style>
</head>
<body>{html_body}</body>
</html>"""

        return HTML(string=html).write_pdf()
    except Exception as e:
        raise RuntimeError(f"ไม่สามารถสร้าง PDF ได้: {e}")
